from django.contrib.humanize.templatetags.humanize import intcomma
from django.utils.html import mark_safe
from provernalog.models import City, ValueFactor
from decimal import Decimal
from docx import Document
import requests
from io import BytesIO

class ParcelMainInformation:
    def __init__(self, parcel):
        self.parcel = parcel
        self.city = City.objects.get(region=self.parcel.region)
        self.current = 'Текущие характеристики' if not self.parcel.approved_cost else "Предыдущие характеристики"
        self.new = 'Новые характеристики' if not self.parcel.approved_cost else 'Текущие характеристики'

    def current_specification(self):
        _parcel = {}
        _parcel["Кадастровый номер"] = self.parcel.cadastral_number
        _parcel["Местоположение объекта"] = self.parcel.address
        _parcel['Площадь объекта (кв.м.)'] = intcomma(self.parcel.area)
        _parcel['Использование объекта (по документу)'] = self.parcel.bydoc
        _parcel['Вид разрешенного использования (по справочнику)'] = self.parcel.utilization.name if self.parcel.utilization else None
        _parcel["Текущая кадастровая стоимость"] = self.parcel.current_cost
        spec_cost = self.parcel.spec_cost(self.parcel.current_cost)
        _parcel['Действующий удельный показатель кадастровой стоимости'] = spec_cost
        if self.parcel.approved_cost:
            _parcel["Предыдущая кадастровая стоимость"] = _parcel["Текущая кадастровая стоимость"]
            _parcel["Предыдущий удельный показатель кадастровой стоимости"] = _parcel['Действующий удельный показатель кадастровой стоимости']
        for item, value in _parcel.items():
            if isinstance(value, Decimal):
                value = intcomma(str(value)) + " руб."
            yield item, value
    current_specification.label = "Текущие характеристики"


    def new_specification(self):
        _parcel = {}
        if self.parcel.approved_cost:
            _parcel['Утвержденная кадастровая стоимость'] = self.parcel.approved_cost
            spec_cost = self.parcel.spec_cost(self.parcel.approved_cost)
            _parcel['Удельный показатель утвержденной кадастровой стоимости'] = spec_cost
            percent_change = self.parcel.percent_change(self.parcel.current_cost, self.parcel.approved_cost)
            _parcel['Относительное изменение утвержденной кадастровой стоимости'] = percent_change
        else:
            _parcel['Кадастровая стоимость из проекта отчета'] = self.parcel.cost_intermediate
            spec_cost = self.parcel.spec_cost(self.parcel.cost_intermediate)
            _parcel['Удельный показатель кадастровой стоимости из проекта отчета'] = spec_cost
            percent_change = self.parcel.percent_change(self.parcel.current_cost, self.parcel.cost_intermediate)
            _parcel['Относительное изменение кадастровой стоимости'] = percent_change
        _parcel['Информация актуальна по состоянию на'] = self.parcel.date_inform
        _parcel['Адрес сайта ГБУ'] = mark_safe(f'<a href="{self.city.source_url}">{self.city.source_url}</a>')
        for item, value in _parcel.items():
            if isinstance(value, Decimal):
                value = intcomma(str(value)) + " руб."
            yield item, value
    new_specification.label = "Новые характеристики"

    def factors(self):
        factors = ValueFactor.objects.filter(parcel=self.parcel)
        for factor in factors:
            item = factor.evaluative_factor.name
            qualitative_value = factor.qualitative_value.strip()
            dimension = factor.evaluative_factor.quantitative_dimension.strip()
            value = f'''{qualitative_value} {dimension}'''
            yield item, value

    def rosreestr(self):
        _parcel = {}
        cadastral_number_array = self.parcel.cadastral_number.split(":")
        cad_num = str(int(cadastral_number_array[0])) + ":" + str(int(cadastral_number_array[1])) + ":" + str(
            int(cadastral_number_array[2])) + ":" + str(int(cadastral_number_array[3]))
        url = f"https://rosreestr.ru/api/online/fir_object/{cad_num}"
        try:
            response = requests.get(url, timeout=2)
            if response.status_code == 200:
                data = response.json()
                object_data = data.get('objectData')
                parcel_data = data.get('parcelData')
                if object_data:
                    _parcel['Наименование'] = object_data.get("objectName")
                    _parcel['Адрес'] = object_data.get('addressNote')
                    object_address = object_data.get('objectAddress')
                    if object_address:
                        _parcel['Кладр'] = object_address.get('kladr')
                if parcel_data:
                    _parcel['Площадь'] = parcel_data.get("areaValue")
                    _parcel['Кадастровая стоимость'] = parcel_data.get('cadCost')
                    _parcel['Дата постановки на учёт'] = parcel_data.get("dateCreate")
                    _parcel['Категория земель'] = parcel_data.get('categoryTypeValue')
                    _parcel['ВРИ'] = parcel_data.get('utilCodeDesc')
                    _parcel['Вид разрешенного использования (по справочнику)'] = parcel_data.get('utilByDoc')
                    _parcel['Тип ОКС'] = parcel_data.get('oksType')
                    _parcel['Этажность'] = parcel_data.get('oksFloors')
                    _parcel['Материал стен'] = parcel_data.get('oksElementsConstruct')
                    _parcel['Год постройки'] = parcel_data.get('oksYearBuilt')
                    _parcel['Год ввода в эксплуатацию'] = parcel_data.get("oksYearUsed")
                    _parcel['Дата утверждения КС'] = parcel_data.get('dateCost')
        except requests.exceptions.ReadTimeout:
            pass
        for item, value in _parcel.items():
            if value:
                yield item, value


class RequestWord:

    def __init__(self, parcel, subject):
        self.city = City.objects.get(region=parcel.region)
        if not self.city.file:
            raise ValueError
        self.document = Document(self.city.file)
        self.parcel = parcel
        self.subject = subject

    def create_word_file(self):
        paragraphs = self.document.paragraphs
        arguments = {}
        group_appraise = self.parcel.group_appraise
        group_type = 'неизвестно'
        if group_appraise:
            if group_appraise.type == 'ZU':
                group_type = 'земельный участок'
            elif group_appraise.type == 'OKS':
                group_type = 'Объект капитального строительства'
        arguments['group_type'] = group_type
        arguments['parcel'] = self.parcel
        arguments['subject'] = self.subject
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        paragraph.text = text.format(**arguments)
        for paragraph in paragraphs:
            text = paragraph.text
            paragraph.text = text.format(**arguments)
        target_stream = BytesIO()
        self.document.save(target_stream)
        target_stream.getvalue()
        target_stream.seek(0)
        return target_stream
